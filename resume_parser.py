"""
Resume Parser Module
Extracts text from PDF/DOCX files and identifies skills using keyword matching.
"""

import re
from io import BytesIO

# PDF parsing
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

# Word document parsing
try:
    from docx import Document
except ImportError:
    Document = None

# Image OCR parsing
try:
    import pytesseract
    from PIL import Image
    ocr_available = True
except ImportError:
    ocr_available = False
    pytesseract = None
    Image = None

# AI Engine for intelligent extraction
try:
    from ollama_engine import chat_with_ollama, clean_response, is_ollama_available
    ollama_module_available = True
except ImportError:
    ollama_module_available = False
    print("Warning: ollama_engine not found, using regex only")


# Comprehensive skill database organized by category
SKILL_DATABASE = {
    # Programming Languages
    "programming_languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "php",
        "swift", "kotlin", "go", "golang", "rust", "scala", "perl", "r", "matlab",
        "objective-c", "dart", "lua", "haskell", "clojure", "elixir", "erlang",
        "fortran", "cobol", "assembly", "bash", "shell", "powershell", "vba"
    ],
    
    # Web Development
    "web_development": [
        "html", "html5", "css", "css3", "sass", "scss", "less", "tailwind",
        "bootstrap", "jquery", "react", "reactjs", "react.js", "angular", "angularjs",
        "vue", "vuejs", "vue.js", "svelte", "next.js", "nextjs", "nuxt", "gatsby",
        "webpack", "vite", "rollup", "parcel", "babel", "npm", "yarn", "pnpm"
    ],
    
    # Backend & Frameworks
    "backend_frameworks": [
        "node.js", "nodejs", "express", "express.js", "django", "flask", "fastapi",
        "spring", "spring boot", "springboot", ".net", "asp.net", "rails", "ruby on rails",
        "laravel", "symfony", "codeigniter", "nest.js", "nestjs", "koa", "hapi",
        "gin", "echo", "fiber", "actix", "rocket"
    ],
    
    # Databases
    "databases": [
        "sql", "mysql", "postgresql", "postgres", "mongodb", "redis", "elasticsearch",
        "cassandra", "dynamodb", "firebase", "sqlite", "oracle", "sql server", "mssql",
        "mariadb", "couchdb", "neo4j", "graphql", "prisma", "sequelize", "mongoose",
        "typeorm", "knex", "drizzle"
    ],
    
    # Cloud & DevOps
    "cloud_devops": [
        "aws", "amazon web services", "azure", "microsoft azure", "gcp", "google cloud",
        "docker", "kubernetes", "k8s", "jenkins", "gitlab ci", "github actions",
        "terraform", "ansible", "puppet", "chef", "vagrant", "cloudformation",
        "ec2", "s3", "lambda", "ecs", "eks", "fargate", "rds", "cloudfront",
        "nginx", "apache", "load balancer", "cdn", "ci/cd", "cicd", "devops"
    ],
    
    # Data Science & AI/ML
    "data_science_ml": [
        "machine learning", "deep learning", "artificial intelligence", "ai", "ml",
        "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "pandas", "numpy",
        "scipy", "matplotlib", "seaborn", "plotly", "jupyter", "notebook",
        "nlp", "natural language processing", "computer vision", "opencv",
        "neural network", "cnn", "rnn", "lstm", "transformer", "bert", "gpt",
        "data mining", "data analysis", "data visualization", "big data",
        "spark", "hadoop", "hive", "kafka", "airflow", "dbt", "snowflake",
        "tableau", "power bi", "looker", "metabase"
    ],
    
    # Mobile Development
    "mobile_development": [
        "android", "ios", "react native", "flutter", "xamarin", "ionic",
        "swiftui", "uikit", "jetpack compose", "kotlin multiplatform",
        "cordova", "phonegap", "expo", "mobile development", "app development"
    ],
    
    # Version Control & Tools
    "tools": [
        "git", "github", "gitlab", "bitbucket", "svn", "mercurial",
        "jira", "confluence", "trello", "asana", "notion", "slack",
        "vs code", "visual studio", "intellij", "pycharm", "eclipse",
        "xcode", "android studio", "postman", "insomnia", "swagger"
    ],
    
    # Testing
    "testing": [
        "unit testing", "integration testing", "e2e testing", "end-to-end",
        "jest", "mocha", "chai", "jasmine", "cypress", "selenium", "playwright",
        "pytest", "unittest", "rspec", "junit", "testng", "cucumber",
        "tdd", "bdd", "test driven", "qa", "quality assurance"
    ],
    
    # Soft Skills
    "soft_skills": [
        "leadership", "teamwork", "communication", "problem solving", "problem-solving",
        "critical thinking", "time management", "project management", "agile", "scrum",
        "kanban", "waterfall", "collaboration", "presentation", "negotiation",
        "mentoring", "coaching", "strategic planning", "decision making",
        "analytical", "creative", "innovative", "detail-oriented", "self-motivated"
    ],
    
    # Business & Management
    "business": [
        "product management", "product owner", "business analysis", "business intelligence",
        "stakeholder management", "budget management", "vendor management",
        "crm", "salesforce", "hubspot", "sap", "erp", "financial analysis",
        "marketing", "seo", "digital marketing", "content marketing", "analytics",
        "excel", "microsoft office", "google workspace", "powerpoint", "word"
    ],
    
    # Security
    "security": [
        "cybersecurity", "information security", "network security", "penetration testing",
        "ethical hacking", "vulnerability assessment", "siem", "soc", "firewall",
        "encryption", "ssl", "tls", "oauth", "jwt", "authentication", "authorization",
        "owasp", "security audit", "compliance", "gdpr", "hipaa", "pci-dss"
    ],
    
    # Design
    "design": [
        "ui", "ux", "ui/ux", "user interface", "user experience", "figma", "sketch",
        "adobe xd", "photoshop", "illustrator", "indesign", "after effects",
        "wireframing", "prototyping", "responsive design", "accessibility",
        "design thinking", "interaction design", "visual design"
    ]
}


def extract_text_from_pdf(file_bytes):
    """
    Extract text content from a PDF file.
    
    Args:
        file_bytes: Bytes content of the PDF file
        
    Returns:
        str: Extracted text from all pages
    """
    if PdfReader is None:
        raise ImportError("PyPDF2 is not installed. Run: pip install PyPDF2")
    
    try:
        pdf_file = BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        
        text_content = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
        
        return "\n".join(text_content)
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {str(e)}")


def extract_text_from_docx(file_bytes):
    """
    Extract text content from a DOCX file.
    
    Args:
        file_bytes: Bytes content of the DOCX file
        
    Returns:
        str: Extracted text from all paragraphs
    """
    if Document is None:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")
    
    try:
        docx_file = BytesIO(file_bytes)
        doc = Document(docx_file)
        
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        return "\n".join(text_content)
    except Exception as e:
        raise Exception(f"Error extracting DOCX text: {str(e)}")


def extract_text_from_image(file_bytes):
    """
    Extract text content from an image file using OCR.
    
    Args:
        file_bytes: Bytes content of the image file
        
    Returns:
        str: Extracted text from the image
    """
    if not ocr_available:
        raise ImportError("pytesseract and Pillow are not installed. Run: pip install pytesseract Pillow")
    
    try:
        image_file = BytesIO(file_bytes)
        image = Image.open(image_file)
        
        # Convert to RGB if necessary (e.g., for PNG with transparency)
        if image.mode in ('RGBA', 'P'):
            image = image.convert('RGB')
        
        # Use Tesseract to extract text
        text = pytesseract.image_to_string(image, lang='eng')
        
        return text
    except Exception as e:
        raise Exception(f"Error extracting text from image: {str(e)}")


def extract_personal_info_llm(text):
    """
    Extract personal information using Ollama LLM.
    Returns dict with extracted fields or None if failed.
    """
    if not ollama_module_available or not is_ollama_available():
        return None
        
    print("🧠 Using LLM for smart resume extraction...")
    
    # Simplified system prompt for faster processing
    system_prompt = """Extract resume info as JSON only. No explanations."""
    
    # Optimized prompt - shorter, clearer, with examples
    # Use /no_think to skip reasoning and get direct output (much faster!)
    # Reduced to 1200 chars - personal info is usually at the top
    prompt = f"""/no_think
Extract from resume:
{text[:1200]}

JSON only:
{{"name":"Full Name","email":"a@b.com","phone":"123","location":"City","bio":"Summary","experience":5}}"""
    
    response, success = chat_with_ollama(prompt, system_prompt, temperature=0.1)
    
    if not success:
        print(f"❌ Ollama Chat Failed: {response}")
        return None
        
    if success:
        try:
            # Debug: Log raw response
            print(f"🔍 LLM raw response (first 200 chars): {response[:200] if response else 'EMPTY'}...", flush=True)
            
            # Clean response to ensure valid JSON
            import json
            cleaned_json = clean_response(response)
            # Find JSON block if embedded
            if '{' in cleaned_json:
                start = cleaned_json.find('{')
                end = cleaned_json.rfind('}') + 1
                cleaned_json = cleaned_json[start:end]
            
            data = json.loads(cleaned_json)
            
            # Validate core fields
            name = data.get('name')
            if not name or len(name) < 2 or name.lower() in ['full name', 'your name', 'name', '[full name]']:
                return None
                
            return {
                'name': name,
                'email': data.get('email'),
                'phone': data.get('phone'),
                'location': data.get('location'),
                'bio': data.get('bio'),
                'experience': data.get('experience')
            }
        except Exception as e:
            print(f"⚠️ LLM JSON extraction failed: {e}")
            return None
    return None

def extract_personal_info(text):
    """
    Extract personal information from resume text.
    Tries LLM first, attempts to fill missing fields with Regex.
    """
    result = {
        'name': None,
        'email': None,
        'phone': None,
        'location': None,
        'bio': None,
        'experience': None
    }
    
    # 1. Try LLM Extraction First
    llm_result = extract_personal_info_llm(text)
    if llm_result:
        print(f"✅ LLM Extraction Success: {llm_result['name']}", flush=True)
        result.update({k: v for k, v in llm_result.items() if v})
        result['debug_info'] = "LLM Used"
    else:
        print("⚠️ LLM Extraction failed or unavailable. Using Regex fallback.", flush=True)
        result['debug_info'] = "Regex Fallback"
    
    # 2. Regex Fallback / Augmentation (fill missing fields)
    
    # Debug: Print first 500 chars of extracted text
    print(f"📄 Extracted text preview: {text[:500]}...", flush=True)

    # Save to debug file
    try:
        import os
        log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
        os.makedirs(log_dir, exist_ok=True)
        with open(os.path.join(log_dir, 'resume_debug.txt'), 'w') as f:
            f.write(text)
    except Exception as e:
        print(f"Failed to write debug log: {e}")
    
    # Extract email using regex (multiple patterns for better accuracy)
    if not result['email']:
        # Standard email pattern
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        email_matches = re.findall(email_pattern, text)
        if not email_matches:
            # Try removing all spaces (fixes bad PDF parsing where characters are spaced out)
            text_no_spaces = re.sub(r'\s+', '', text)
            email_matches = re.findall(email_pattern, text_no_spaces)
            
        if email_matches:
            result['email'] = email_matches[0].lower()
        else:
            # Try lenient pattern for truncated/OCR emails (e.g. "name@gm" from garbled PDFs)
            lenient_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+'
            lenient_matches = re.findall(lenient_pattern, text)
            for match in lenient_matches:
                # Try to reconstruct common domains
                if '@gm' in match.lower() or '@gmail' in match.lower():
                    result['email'] = match.split('@')[0] + '@gmail.com'
                    break
                elif '@yahoo' in match.lower() or '@yah' in match.lower():
                    result['email'] = match.split('@')[0] + '@yahoo.com'
                    break
                elif '@hotm' in match.lower() or '@outlook' in match.lower():
                    result['email'] = match.split('@')[0] + '@outlook.com'
                    break
                elif '@em' in match.lower() or '@email' in match.lower():
                    result['email'] = match.split('@')[0] + '@email.com'
                    break
                elif len(match) > 5:  # Just use as-is if we found something
                    result['email'] = match.lower()
                    break
    
    # Extract phone number using regex (various formats)
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # +60 123 456 7890
        r'\+?\d{1,3}[-.\s]?\d{2,4}[-.\s]?\d{3,4}[-.\s]?\d{3,4}',  # Various formats
        r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}',  # (012) 345-6789
        r'\d{10,12}'  # Plain 10-12 digit number
    ]
    for pattern in phone_patterns:
        phone_matches = re.findall(pattern, text)
        if phone_matches:
            # Clean up the phone number
            phone = re.sub(r'[^\d+]', '', phone_matches[0])
            if len(phone) >= 10:
                result['phone'] = phone_matches[0]
                break
    
    # Extract location - look for city, state, country patterns
    location_patterns = [
        # "New York, USA" or "New York, NY" or "Kuala Lumpur, Malaysia"
        r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*,\s*([A-Z]{2,3}|[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b',
        # After "Address" or "Location" label
        r'(?:address|location)[:\s]+([A-Za-z\s,]+(?:USA|Malaysia|Singapore|UK|Australia|India|Canada))',
    ]
    
    # Common countries/states to validate location
    location_keywords = ['usa', 'malaysia', 'singapore', 'uk', 'australia', 'india', 'canada',
                        'new york', 'california', 'texas', 'kuala lumpur', 'london', 'sydney']
    
    text_lower = text.lower()
    for pattern in location_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                loc = ', '.join(match).strip()
            else:
                loc = match.strip()
            # Validate it looks like a location
            if any(kw in loc.lower() for kw in location_keywords):
                result['location'] = loc
                break
        if result['location']:
            break
    
    # If no location found via patterns, look for location keywords in text
    if not result['location']:
        for kw in location_keywords:
            if kw in text_lower:
                # Try to extract the surrounding context
                match = re.search(rf'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s*,?\s*{kw}', text, re.IGNORECASE)
                if match:
                    result['location'] = match.group(0).strip()
                    break
    
    # --- IMPROVED NAME EXTRACTION ---
    # Strategy: Look for all-caps lines that might be the name (merging if split)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    cleaned_lines = [] 
    for line in lines:
        cleaned_lines.append(re.sub(r'[^a-zA-Z\s]', '', line).strip())

    headers = {'WORK EXPERIENCE', 'EDUCATION', 'SKILLS', 'CERTIFICATIONS', 'PROFILE', 'SUMMARY', 'CONTACT', 'LANGUAGES', 'PROJECTS', 'ACHIEVEMENTS', 'AWARDS', 'PUBLICATIONS', 'TECHNICAL', 'OBJECTIVE', 'EXPERTISE', 'COMPETENCIES'}
    skip_words = ['resume', 'cv', 'curriculum', 'vitae', 'summary', 'objective', 
                  'experience', 'education', 'skills', 'contact', 'profile', 'about',
                  'address', 'phone', 'email', 'mobile', 'tel', 'linkedin', 'github',
                  'portfolio', 'website', 'professional', 'career', 'personal', 'info',
                  'information', 'details', 'page', 'date', 'nationality', 'gender',
                  'present', 'current', 'location', 'city', 'state', 'country',
                  'work', 'language', 'languages', 'hobbies', 'hobby', 'interests',
                  'references', 'reference', 'designer', 'developer', 'engineer',
                  'manager', 'analyst', 'specialist', 'consultant', 'senior', 'junior',
                  'new', 'york', 'usa', 'texas', 'california', 'bachelor', 'master',
                  'figma', 'adobe', 'sketch', 'invision', 'photoshop', 'illinois', 'chicago',
                  # Additional job titles
                  'intern', 'executive', 'director', 'lead', 'head', 'associate',
                  'coordinator', 'administrator', 'assistant', 'officer', 'supervisor',
                  'technician', 'programmer', 'architect', 'scientist', 'researcher',
                  'accountant', 'marketing', 'sales', 'support', 'customer', 'service',
                  'student', 'graduate', 'fresher', 'trainee', 'apprentice', 'learner',
                  # Malaysian locations and address terms
                  'kuala', 'lumpur', 'selangor', 'penang', 'johor', 'pahang', 'perak',
                  'kedah', 'kelantan', 'terengganu', 'sabah', 'sarawak', 'melaka',
                  'negeri', 'sembilan', 'perlis', 'putrajaya', 'labuan', 'malaysia',
                  'jalan', 'taman', 'cheras', 'bangsar', 'damansara', 'ampang', 'subang',
                  'petaling', 'puchong', 'shah', 'alam', 'klang', 'lenseng', 'ipoh',
                  # Other locations
                  'singapore', 'london', 'sydney', 'australia', 'india', 'mumbai',
                  'bangalore', 'delhi', 'chennai', 'dubai', 'hong', 'kong', 'china',
                  'anywhere', 'city', 'street', 'avenue', 'road', 'lane',
                  # Common resume terms
                  'years', 'year', 'months', 'month', 'january', 'february', 'march',
                  'april', 'may', 'june', 'july', 'august', 'september', 'october',
                  'november', 'december', 'jan', 'feb', 'mar', 'apr', 'jun', 'jul',
                  'aug', 'sep', 'oct', 'nov', 'dec', 'profile', 'contact',
                  # Tech terms that might appear prominently
                  'python', 'java', 'javascript', 'react', 'node', 'angular', 'vue',
                  'html', 'css', 'sql', 'mongodb', 'mysql', 'aws', 'azure', 'docker',
                  # Company types
                  'company', 'corporation', 'enterprise', 'solutions', 'technologies',
                  'systems', 'software', 'consulting', 'agency', 'group', 'limited', 'ltd',
                  'inc', 'llc', 'pvt', 'private', 'sdn', 'bhd', 'berhad',
                  # Languages (people often list these prominently)
                  'chinese', 'english', 'malay', 'mandarin', 'cantonese', 'tamil', 'hindi',
                  'spanish', 'french', 'german', 'japanese', 'korean', 'arabic', 'russian',
                  'native', 'fluent', 'basic', 'intermediate', 'advanced', 'proficient',
                  # Common resume section headers
                  'languages', 'language', 'skills', 'skill', 'education', 'experience',
                  'work', 'volunteer', 'references', 'certifications', 'achievements',
                  'hobbies', 'interests', 'objective', 'summary', 'degree', 'institute',
                  # OCR artifacts
                  'hello', 'reallygreatsite', 'linkedin', 'facebook', 'twitter']

    # PRIORITY 1: Look for ALL CAPS names (very common in resumes)
    # These are usually 2-3 words, all uppercase, like "EDDEE LEE" or "CHONG Y.L."
    all_caps_name = None
    for line in lines[:30]:  # Check first 30 lines
        line_clean = line.strip()
        
        # Skip lines that contain email indicators, numbers, or special chars
        if '@' in line_clean or any(c.isdigit() for c in line_clean):
            continue
        if '(' in line_clean or ')' in line_clean or '[' in line_clean:
            continue
        if 'http' in line_clean.lower() or 'www.' in line_clean.lower():
            continue
        
        # Check if line is mostly uppercase and looks like a name
        if len(line_clean) > 3 and len(line_clean) < 40:
            words = line_clean.split()
            if 1 <= len(words) <= 4:
                # Check if it's mostly uppercase (allowing for initials like Y.L.)
                upper_chars = sum(1 for c in line_clean if c.isupper() or c in '.')
                total_alpha = sum(1 for c in line_clean if c.isalpha() or c in '.')
                if total_alpha > 2 and upper_chars / total_alpha > 0.6:
                    # Make sure none of the words are skip words
                    word_lower = [w.lower().replace('.', '') for w in words]
                    if not any(w in skip_words for w in word_lower):
                        # This looks like an ALL CAPS name
                        all_caps_name = line_clean.title()
                        # Handle initials like Y.L.
                        all_caps_name = re.sub(r'([A-Z])\. ?([A-Z])\.?', r'\1.\2.', all_caps_name)
                        print(f"🎯 ALL CAPS name detected: {all_caps_name}", flush=True)
                        break
    
    # Use ALL CAPS name if found and LLM didn't set one
    if all_caps_name and not result['name']:
        result['name'] = all_caps_name
    
    # PRIORITY 2: Check the very first non-empty line (most resumes have name at top)
    first_valid_name = None
    if not result['name']:  # Only if we don't have a name yet
        for i, line in enumerate(lines[:5]):  # Check first 5 lines
            clean = re.sub(r'[^a-zA-Z\s]', '', line).strip()
            words = clean.split()
            
            # Skip empty or single character lines
            if len(words) < 2:
                continue
            
            # Skip if contains skip words
            if any(w.lower() in skip_words for w in words):
                continue
            
            # Valid name: 2-4 words, looks like a person's name
            if 2 <= len(words) <= 4:
                # Check if all words are alphabetic
                all_alpha = all(w.replace('.', '').replace('-', '').replace("'", '').isalpha() for w in words)
                if all_alpha:
                    # This is likely the name!
                    first_valid_name = ' '.join(w.title() if w.isupper() or w.islower() else w for w in words)
                    print(f"🎯 First-line name detected: {first_valid_name}", flush=True)
                    break
    
    # Only use first-line name if LLM didn't set the name
    if first_valid_name and not result['name']:
        result['name'] = first_valid_name
    elif not result['name']:
        name_candidates = []
        # Fallback: Scan first 30 lines for name candidates
        for i, line in enumerate(lines[:30]):
            clean_text = cleaned_lines[i]
            
            # Check conditions
            is_upper = line.isupper()
            is_title = line.istitle()
            is_header = line.upper() in headers or any(h in line.upper() for h in headers)
            
            # Valid name candidate?
            if (is_upper or is_title) and not is_header and 2 < len(clean_text) < 30:
                # Filter out skip words
                words = clean_text.split()
                if not any(w.lower() in skip_words for w in words):
                     name_candidates.append({'index': i, 'text': clean_text, 'original': line})

        # Try to find adjacent candidates to merge (e.g. "DENNIS" \n "SCHERRER")
        extracted_name = None
        
        if name_candidates:
            for i in range(len(name_candidates) - 1):
                curr = name_candidates[i]
                next_cand = name_candidates[i+1]
                if next_cand['index'] == curr['index'] + 1:
                    combined = f"{curr['text']} {next_cand['text']}"
                    if len(combined.split()) >= 2:
                        extracted_name = combined.title()
                        break
            
            # If no merge, take first valid looking candidate
            if not extracted_name:
                for cand in name_candidates:
                    if len(cand['text'].split()) >= 2:
                        extracted_name = cand['text'].title()
                        break

        if extracted_name:
            result['name'] = extracted_name
            print(f"🔍 Extracted Name (Smart Logic): {result['name']}", flush=True)


    # --- IMPROVED LOCATION EXTRACTION (Artifact Handling) ---
    if not result['location']:
        # Look for "Q Houston, TX" pattern or similar
        loc_pattern = r'(?i)(?:^|\n|[\s\|•Q])\s*([A-Za-z\s]+,[ \t]*[A-Z]{2,})'
        loc_matches = re.findall(loc_pattern, text)
        for loc in loc_matches:
             # Basic validation
             if any(kw in loc.lower() for kw in ['houston', 'york', 'kuala', 'lumpur', 'selangor', 'penang', 'johor', 'texas', 'california']):
                 result['location'] = loc.strip()
                 break
    
    # Second pass: If no all-caps name found, look for regular names
    if result['name'] is None:
        for line in lines[:20]:  # Increased scan range
            line = line.strip()
            # Skip empty lines
            if not line or len(line) < 3:
                continue
            # Skip lines with email/phone/url
            if '@' in line or 'http' in line.lower() or 'www.' in line.lower():
                continue
            # Skip lines that are too long (likely descriptions)
            if len(line) > 40:
                continue
            # Skip lines with numbers
            if any(c.isdigit() for c in line):
                continue
            # Skip lines that look like headers/sections
            line_lower = line.lower()
            if any(word == line_lower.strip() for word in skip_words):
                continue
            # Skip lines with special characters
            if re.search(r'[|•·:;,@#$%^&*()+=\[\]{}]', line):
                continue
            
            # Check if line looks like a name
            words = line.split()
            if 2 <= len(words) <= 4:
                # Check if all words are alphabetic
                valid_name = True
                for w in words:
                    clean_word = w.replace('.', '').replace('-', '').replace("'", '')
                    if not clean_word.isalpha():
                        valid_name = False
                        break
                    # Skip if word is in skip list
                    if w.lower() in skip_words:
                        valid_name = False
                        break
                
                if valid_name:
                    result['name'] = ' '.join(w.title() if w.isupper() or w.islower() else w for w in words)
                    break
                    
    # FALLBACK: If still no name, just take the first non-empty line that isn't a skip word
    if result['name'] is None:
        for line in lines[:5]:
            line = line.strip()
            if line and len(line) > 3 and not any(w in line.lower() for w in skip_words) and not '@' in line:
                 words = line.split()
                 if 2 <= len(words) <= 4:
                     result['name'] = line.title()
                     break

    print(f"🔍 Extracted Name: {result['name']}")  # Debug print
    
    # Extract experience years
    text_lower = text.lower()
    
    # Method 1: Look for explicit experience mentions
    exp_patterns = [
        r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)',  # "5 years experience" or "5+ years of experience"
        r'experience[:\s]+(\d+)\+?\s*years?',  # "Experience: 5 years"
        r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:work|professional|industry)',  # "5 years of work"
        r'over\s*(\d+)\+?\s*years?',  # "over 5 years"
        r'more\s*than\s*(\d+)\+?\s*years?',  # "more than 5 years"
    ]
    
    for pattern in exp_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                years = int(match.group(1))
                if 0 < years <= 50:  # Reasonable range
                    result['experience'] = years
                    break
            except:
                pass
    
    # Method 2: If no explicit mention, try to calculate from work history dates
    if result['experience'] is None:
        # Look for year ranges in various formats:
        # "Jan 2020 - Dec 2022", "January 2018 - December 2019", "2019 - 2023", "2019 - Present"
        year_patterns = [
            # Month Year - Month Year format (e.g., "Jan 2020 - Dec 2022")
            r'(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s*(20\d{2})\s*[-–—to]+\s*(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s*(20\d{2}|present|current|now)',
            # Just year ranges: "2019 - 2023" or "2019 - Present"
            r'(20\d{2})\s*[-–—to]+\s*(20\d{2}|present|current|now)',
            r'(19\d{2})\s*[-–—to]+\s*(20\d{2}|present|current|now)',
        ]
        
        years_worked = set()
        current_year = 2026  # Use current year
        
        for pattern in year_patterns:
            matches = re.findall(pattern, text_lower)
            for match_tuple in matches:
                try:
                    # Get start and end year (last two elements in tuple)
                    if len(match_tuple) >= 2:
                        start_year = match_tuple[-2] if match_tuple[-2].isdigit() else match_tuple[0]
                        end_year = match_tuple[-1]
                    else:
                        continue
                    
                    start = int(start_year)
                    if end_year in ['present', 'current', 'now']:
                        end = current_year
                    else:
                        end = int(end_year)
                    
                    # Calculate years for this job period
                    if start <= end and start >= 1970 and end <= current_year + 1:
                        for y in range(start, end + 1):
                            years_worked.add(y)
                except:
                    pass
        
        if years_worked:
            # Calculate experience as (latest year - earliest year) instead of counting individual years
            # This avoids including education years in the count
            min_year = min(years_worked)
            max_year = max(years_worked)
            
            # If the max year is current or recent, use current year
            if max_year >= current_year - 1:
                total_years = current_year - min_year
            else:
                total_years = max_year - min_year + 1
            
            if 0 < total_years <= 50:
                result['experience'] = total_years
    
    # If experience seems too high (includes education years), recalculate
    if result['experience'] and result['experience'] > 15:
        # Find all year ranges in the text with their positions
        year_pattern = r'(20\d{2}|19\d{2})\s*[-–—to]+\s*(20\d{2}|present|current|now)'
        
        # Find education section position to exclude its dates
        edu_match = re.search(r'\beducation\b', text_lower)
        edu_pos = edu_match.start() if edu_match else len(text_lower)
        
        # Find work dates - those NOT near the education section
        work_years = set()
        for match in re.finditer(year_pattern, text_lower, re.IGNORECASE):
            match_pos = match.start()
            
            # Skip dates that are within 200 chars after "education" header
            if edu_match and edu_pos < match_pos < edu_pos + 200:
                continue
            
            try:
                start_year = int(match.group(1))
                end_year_str = match.group(2).lower()
                end_year = current_year if end_year_str in ['present', 'current', 'now'] else int(end_year_str)
                
                if 1990 <= start_year <= current_year and start_year <= end_year:
                    work_years.add(start_year)
                    work_years.add(end_year)
            except:
                pass
        
        if work_years:
            min_work = min(work_years)
            max_work = max(work_years)
            if max_work >= current_year - 1:
                result['experience'] = current_year - min_work
            else:
                result['experience'] = max_work - min_work + 1
    
    # Extract bio/summary/objective - improved extraction
    # Look for Profile/Summary sections and get the content after the header
    bio_headers = [
        r'\bprofile\b\s*\n',
        r'\bsummary\b\s*\n', 
        r'\bobjective\b\s*\n',
        r'\babout\s*me\b\s*\n',
        r'\bprofessional\s*summary\b\s*\n',
        r'\bcareer\s*objective\b\s*\n',
    ]
    
    for header_pattern in bio_headers:
        match = re.search(header_pattern, text_lower)
        if match:
            # Get text after the header until next section
            start_pos = match.end()
            # Find next section header
            next_section = re.search(r'\n(?:experience|education|skills|work|employment|languages|hobbies|contact)\b', text_lower[start_pos:], re.IGNORECASE)
            if next_section:
                end_pos = start_pos + next_section.start()
            else:
                end_pos = min(start_pos + 500, len(text))
            
            bio_text = text[start_pos:end_pos].strip()
            # Clean up
            bio_text = re.sub(r'\s+', ' ', bio_text)
            bio_text = bio_text.strip()
            
            # Make sure we got meaningful content
            if len(bio_text) > 30 and not bio_text.lower().startswith(('jan', 'feb', 'mar', 'apr', 'may', 'jun', 'jul', 'aug', 'sep', 'oct', 'nov', 'dec')):
                result['bio'] = bio_text[:500]
                break
    
    # Auto-generate bio if none found from resume
    if not result['bio'] and result['name']:
        # Create a simple bio from available info
        name = result['name']
        exp = result.get('experience')
        
        # Try to find job title from first lines
        job_title = None
        for line in lines[:10]:
            line_clean = line.strip()
            # Common job title patterns
            if any(title in line_clean.lower() for title in ['engineer', 'developer', 'manager', 'designer', 'analyst', 'specialist', 'consultant', 'coordinator']):
                if len(line_clean) < 50 and not '@' in line_clean:
                    job_title = line_clean
                    break
        
        if job_title and exp:
            result['bio'] = f"Professional with {exp} years of experience as {job_title}."
        elif job_title:
            result['bio'] = f"Experienced professional working as {job_title}."
        elif exp:
            result['bio'] = f"Professional with {exp} years of work experience."
        else:
            result['bio'] = "Experienced professional seeking new opportunities."
    
    return result


def extract_skills(text):
    """
    Extract skills from resume text using keyword matching.
    
    Args:
        text: Plain text content from resume
        
    Returns:
        dict: {
            'skills': list of unique matched skills,
            'skills_by_category': dict of skills grouped by category,
            'total_matches': total count of skill matches
        }
    """
    # Normalize text for matching
    text_lower = text.lower()
    # Replace common separators with spaces for better matching
    text_normalized = re.sub(r'[,;|•·\-–—/\\]', ' ', text_lower)
    text_normalized = re.sub(r'\s+', ' ', text_normalized)
    
    found_skills = set()
    skills_by_category = {}
    
    for category, skills in SKILL_DATABASE.items():
        category_matches = []
        for skill in skills:
            skill_lower = skill.lower()
            # Use word boundary matching for more accuracy
            # Create pattern that handles variations
            pattern = r'\b' + re.escape(skill_lower).replace(r'\ ', r'\s*') + r'\b'
            if re.search(pattern, text_normalized):
                # Capitalize skill for display
                display_skill = skill.title() if len(skill) > 3 else skill.upper()
                # Special cases for acronyms and specific capitalization
                special_caps = {
                    'javascript': 'JavaScript', 'typescript': 'TypeScript',
                    'nodejs': 'Node.js', 'node.js': 'Node.js',
                    'reactjs': 'React.js', 'react.js': 'React.js',
                    'vuejs': 'Vue.js', 'vue.js': 'Vue.js',
                    'angularjs': 'AngularJS', 'mongodb': 'MongoDB',
                    'postgresql': 'PostgreSQL', 'mysql': 'MySQL',
                    'graphql': 'GraphQL', 'nosql': 'NoSQL',
                    'html5': 'HTML5', 'css3': 'CSS3',
                    'aws': 'AWS', 'gcp': 'GCP', 'api': 'API',
                    'ci/cd': 'CI/CD', 'devops': 'DevOps',
                    'github': 'GitHub', 'gitlab': 'GitLab',
                    'pytorch': 'PyTorch', 'tensorflow': 'TensorFlow',
                    'opencv': 'OpenCV', 'scikit-learn': 'Scikit-learn',
                    'ui/ux': 'UI/UX', 'seo': 'SEO',
                    'oauth': 'OAuth', 'jwt': 'JWT',
                }
                display_skill = special_caps.get(skill_lower, display_skill)
                
                found_skills.add(display_skill)
                category_matches.append(display_skill)
        
        if category_matches:
            skills_by_category[category] = list(set(category_matches))
    
    # Sort skills alphabetically
    sorted_skills = sorted(list(found_skills))
    
    return {
        'skills': sorted_skills,
        'skills_by_category': skills_by_category,
        'total_matches': len(sorted_skills)
    }


def parse_resume(file_bytes, filename):
    """
    Main function to parse a resume and extract skills.
    
    Args:
        file_bytes: Bytes content of the file
        filename: Original filename to determine file type
        
    Returns:
        dict: {
            'success': bool,
            'skills': list of extracted skills,
            'skills_by_category': dict of skills by category,
            'raw_text': extracted text (truncated),
            'error': error message if failed
        }
    """
    try:
        # Determine file type
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            text = extract_text_from_pdf(file_bytes)
        elif filename_lower.endswith('.docx'):
            text = extract_text_from_docx(file_bytes)
        elif filename_lower.endswith('.doc'):
            return {
                'success': False,
                'error': 'Old .doc format not supported. Please convert to .docx or .pdf'
            }
        elif filename_lower.endswith(('.jpg', '.jpeg', '.png')):
            if not ocr_available:
                return {
                    'success': False,
                    'error': 'OCR not available. Please install pytesseract and Pillow.'
                }
            text = extract_text_from_image(file_bytes)
        else:
            return {
                'success': False,
                'error': 'Unsupported file format. Please upload PDF, DOCX, JPG, or PNG.'
            }
        
        if not text or len(text.strip()) < 50:
            return {
                'success': False,
                'error': 'Could not extract text from file. The file may be empty or image-based.'
            }
        
        # Validate if the file is actually a resume/CV
        resume_keywords = [
            'experience', 'education', 'skills', 'work', 'employment', 'career',
            'professional', 'qualification', 'certificate', 'degree', 'university',
            'college', 'company', 'position', 'role', 'job', 'responsibilities',
            'achievements', 'project', 'contact', 'email', 'phone', 'address',
            'summary', 'objective', 'profile', 'about', 'expertise', 'competencies',
            'training', 'intern', 'graduate', 'bachelor', 'master', 'phd', 'diploma'
        ]
        
        text_lower = text.lower()
        keyword_matches = sum(1 for kw in resume_keywords if kw in text_lower)
        
        # If less than 3 resume keywords found, likely not a resume
        if keyword_matches < 3:
            return {
                'success': False,
                'error': 'This file does not appear to be a resume/CV. Please upload your actual resume or CV document containing your work experience, education, and skills.'
            }
        
        # Extract personal information (name, email, phone, bio)
        personal_info = extract_personal_info(text)
        
        # Extract skills
        skill_result = extract_skills(text)
        
        if skill_result['total_matches'] == 0:
            return {
                'success': True,
                'skills': [],
                'skills_by_category': {},
                'raw_text': text[:500] + '...' if len(text) > 500 else text,
                'full_text': text,
                'message': 'No skills detected. You may need to add skills manually.',
                'personal_info': personal_info
            }
        
        return {
            'success': True,
            'skills': skill_result['skills'],
            'skills_by_category': skill_result['skills_by_category'],
            'raw_text': text[:500] + '...' if len(text) > 500 else text,
            'full_text': text,
            'total_matches': skill_result['total_matches'],
            'personal_info': personal_info,
            'extraction_method': personal_info.get('debug_info', 'Unknown')
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }
